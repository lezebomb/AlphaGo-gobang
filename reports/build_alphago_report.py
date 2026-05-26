# -*- coding: utf-8 -*-
"""Build a polished 7-8 page DOCX report for the AlphaGo/Gomoku assignment."""

from __future__ import annotations

import json
import math
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
PROJECT = ROOT / "AlphaZero_Gomoku"
RESULTS_PATH = PROJECT / "experiments" / "results_8x8.json"
SWEEP_PATH = PROJECT / "experiments" / "playout_sweep_8x8.json"
ASSET_DIR = ROOT / "reports" / "assets"
OUT_PATH = ROOT / "AlphaGo与围棋AI家族_实验报告_8x8强化版.docx"
FONT_PATH = Path("C:/Windows/Fonts/simhei.ttf")


BLUE = (31, 78, 121)
LIGHT_BLUE = "DDEBF7"
PAPER = (250, 252, 255)
INK = (35, 35, 35)


def pil_font(size, bold=False):
    if FONT_PATH.exists():
        return ImageFont.truetype(str(FONT_PATH), size)
    return ImageFont.load_default()


def draw_centered(draw, box, text, font, fill=INK):
    left, top, right, bottom = box
    bbox = draw.textbbox((0, 0), text, font=font)
    x = left + (right - left - (bbox[2] - bbox[0])) / 2
    y = top + (bottom - top - (bbox[3] - bbox[1])) / 2
    draw.text((x, y), text, font=font, fill=fill)


def build_graph_logic_image(path):
    img = Image.new("RGB", (1180, 640), PAPER)
    draw = ImageDraw.Draw(img)
    title_font = pil_font(34)
    text_font = pil_font(22)
    small_font = pil_font(18)

    draw.text((40, 28), "图论视角：五子棋局面 = 棋盘图上的路径与威胁", font=title_font, fill=BLUE)
    ox, oy, step = 110, 125, 54
    for i in range(8):
        draw.line((ox, oy + i * step, ox + 7 * step, oy + i * step), fill=(145, 155, 165), width=2)
        draw.line((ox + i * step, oy, ox + i * step, oy + 7 * step), fill=(145, 155, 165), width=2)
    for r in range(8):
        for c in range(8):
            x, y = ox + c * step, oy + (7 - r) * step
            draw.ellipse((x - 5, y - 5, x + 5, y + 5), fill=(95, 105, 115))

    black_path = [(3, 4), (4, 3), (5, 2), (2, 5), (1, 6)]
    white_blocks = [(3, 2), (4, 2), (2, 2), (1, 2)]

    def xy(pos):
        r, c = pos
        return ox + c * step, oy + (7 - r) * step

    for a, b in zip(black_path, black_path[1:]):
        draw.line((*xy(a), *xy(b)), fill=(33, 96, 170), width=6)
    for pos in white_blocks:
        x, y = xy(pos)
        draw.ellipse((x - 18, y - 18, x + 18, y + 18), fill=(245, 245, 245), outline=(80, 80, 80), width=3)
    for pos in black_path:
        x, y = xy(pos)
        draw.ellipse((x - 19, y - 19, x + 19, y + 19), fill=(33, 33, 33), outline=(33, 96, 170), width=3)

    legend_x = 610
    draw.rounded_rectangle((legend_x, 120, 1100, 505), radius=18, fill=(236, 243, 250), outline=(176, 198, 220), width=2)
    lines = [
        "棋盘点：交叉点 v ∈ V",
        "棋盘边：横、竖、斜方向相邻关系 e ∈ E",
        "五连胜：同色棋子形成长度 4 的路径",
        "威胁：接近五连的开放路径或双向延展",
        "MCTS 会在这些局部路径与全局胜率间权衡",
    ]
    y = 150
    for line in lines:
        draw.text((legend_x + 28, y), line, font=text_font, fill=INK)
        y += 58
    draw.text((110, 575), "样例终局来自本次 8x8 复现实验：AI 通过斜向路径完成五连。", font=small_font, fill=(90, 90, 90))
    img.save(path)


def build_mcts_image(path):
    img = Image.new("RGB", (1180, 620), PAPER)
    draw = ImageDraw.Draw(img)
    title_font = pil_font(34)
    node_font = pil_font(20)
    small_font = pil_font(18)
    draw.text((40, 28), "AlphaZero 风格 MCTS：在搜索树上迭代改进策略", font=title_font, fill=BLUE)

    nodes = {
        "root": (570, 115, "s0\nN=120"),
        "a": (250, 285, "a1\nP=0.31\nQ=0.42"),
        "b": (570, 285, "a2\nP=0.48\nQ=0.57"),
        "c": (890, 285, "a3\nP=0.11\nQ=0.12"),
        "d": (460, 455, "leaf\nv=0.73"),
        "e": (680, 455, "leaf\nv=0.18"),
    }
    edges = [("root", "a"), ("root", "b"), ("root", "c"), ("b", "d"), ("b", "e")]

    def box(key):
        x, y, label = nodes[key]
        return x - 70, y - 45, x + 70, y + 45

    for s, t in edges:
        sx, sy, _ = nodes[s]
        tx, ty, _ = nodes[t]
        draw.line((sx, sy + 45, tx, ty - 45), fill=(120, 140, 160), width=4)
    for key, (x, y, label) in nodes.items():
        fill = (224, 238, 251) if key != "b" else (211, 232, 214)
        draw.rounded_rectangle(box(key), radius=18, fill=fill, outline=(70, 100, 130), width=3)
        lines = label.split("\n")
        yy = y - 24
        for line in lines:
            draw_centered(draw, (x - 65, yy - 12, x + 65, yy + 18), line, node_font, fill=INK)
            yy += 28

    formula = "选择：argmax_a [ Q(s,a) + c_puct · P(s,a) · sqrt(N(s)) / (1 + N(s,a)) ]"
    draw.rounded_rectangle((85, 520, 1095, 585), radius=14, fill=(245, 248, 252), outline=(190, 205, 220), width=2)
    draw.text((110, 540), formula, font=small_font, fill=(45, 45, 45))
    for txt, pos in [
        ("1 选择", (180, 150)),
        ("2 扩展", (765, 235)),
        ("3 评估", (735, 405)),
        ("4 回传", (420, 350)),
    ]:
        draw.text(pos, txt, font=small_font, fill=(180, 70, 50))
    img.save(path)


def build_eval_image(path, results, sweep):
    img = Image.new("RGB", (1180, 620), PAPER)
    draw = ImageDraw.Draw(img)
    title_font = pil_font(34)
    axis_font = pil_font(20)
    small_font = pil_font(17)
    draw.text((40, 28), "8x8 五子棋模型评估：胜率与搜索预算", font=title_font, fill=BLUE)

    summaries = [
        ("随机", results["matches"]["alpha_vs_random"]["summary"]),
        ("启发式", results["matches"]["alpha_vs_heuristic"]["summary"]),
        ("纯MCTS", results["matches"]["alpha_vs_pure_mcts"]["summary"]),
    ]
    chart_left, chart_top, chart_bottom = 90, 150, 500
    draw.line((chart_left, chart_bottom, 565, chart_bottom), fill=(90, 90, 90), width=2)
    draw.line((chart_left, chart_top, chart_left, chart_bottom), fill=(90, 90, 90), width=2)
    for i in range(6):
        y = chart_bottom - i * 60
        draw.line((chart_left - 6, y, 565, y), fill=(225, 230, 235), width=1)
        draw.text((42, y - 12), "{:.1f}".format(i / 5), font=small_font, fill=(90, 90, 90))
    bar_w = 80
    for i, (name, summary) in enumerate(summaries):
        x = chart_left + 90 + i * 125
        h = int(summary["alpha_score"] * 300)
        draw.rectangle((x, chart_bottom - h, x + bar_w, chart_bottom), fill=(33, 96, 170))
        draw.text((x + 10, chart_bottom - h - 32), "{:.0%}".format(summary["alpha_score"]), font=axis_font, fill=BLUE)
        draw_centered(draw, (x - 15, chart_bottom + 16, x + bar_w + 15, chart_bottom + 45), name, axis_font)
    draw.text((125, 108), "不同对手下的 AlphaZero-MCTS 得分", font=axis_font, fill=INK)

    x0, y0 = 680, 150
    draw.text((x0, 108), "搜索次数敏感性（对启发式玩家）", font=axis_font, fill=INK)
    draw.line((x0, chart_bottom, 1090, chart_bottom), fill=(90, 90, 90), width=2)
    draw.line((x0, chart_top, x0, chart_bottom), fill=(90, 90, 90), width=2)
    settings = sweep["settings"]
    max_time = max(s["summary"]["avg_sec_per_move"] for s in settings)
    points = []
    for i, setting in enumerate(settings):
        px = x0 + 45 + i * 105
        py = chart_bottom - int(setting["summary"]["avg_sec_per_move"] / max_time * 260)
        points.append((px, py))
        draw.ellipse((px - 7, py - 7, px + 7, py + 7), fill=(192, 80, 77))
        draw_centered(draw, (px - 40, chart_bottom + 15, px + 40, chart_bottom + 45), str(setting["alpha_playouts"]), small_font)
        draw.text((px - 28, py - 32), "{:.3f}s".format(setting["summary"]["avg_sec_per_move"]), font=small_font, fill=(120, 60, 55))
    for a, b in zip(points, points[1:]):
        draw.line((*a, *b), fill=(192, 80, 77), width=4)
    draw.text((x0 + 25, 525), "横轴：每步 playouts；纵向：平均每步耗时。胜率均为 100%。", font=small_font, fill=(80, 80, 80))
    img.save(path)


def build_sample_board_image(path, sample_game):
    img = Image.new("RGB", (900, 620), PAPER)
    draw = ImageDraw.Draw(img)
    title_font = pil_font(30)
    small_font = pil_font(18)
    draw.text((40, 28), "样例对局终局：AlphaZero-MCTS 对纯 MCTS", font=title_font, fill=BLUE)
    ox, oy, step = 120, 115, 54
    for i in range(8):
        draw.line((ox, oy + i * step, ox + 7 * step, oy + i * step), fill=(145, 155, 165), width=2)
        draw.line((ox + i * step, oy, ox + i * step, oy + 7 * step), fill=(145, 155, 165), width=2)
        draw.text((ox - 45, oy + i * step - 12), str(7 - i), font=small_font, fill=(85, 85, 85))
        draw.text((ox + i * step - 5, oy + 7 * step + 28), str(i), font=small_font, fill=(85, 85, 85))

    moves = sample_game["transcript"]
    for idx, item in enumerate(moves, 1):
        row, col = item["location"]
        x, y = ox + col * step, oy + (7 - row) * step
        if item["player"] == 1:
            fill, text_fill = (35, 35, 35), (255, 255, 255)
        else:
            fill, text_fill = (246, 246, 246), (40, 40, 40)
        draw.ellipse((x - 21, y - 21, x + 21, y + 21), fill=fill, outline=(50, 80, 120), width=3)
        draw_centered(draw, (x - 18, y - 17, x + 18, y + 17), str(idx), small_font, text_fill)
    draw.rounded_rectangle((585, 125, 845, 405), radius=16, fill=(236, 243, 250), outline=(176, 198, 220), width=2)
    draw.text((610, 155), "评估设置", font=pil_font(23), fill=BLUE)
    info = [
        "棋盘：8x8",
        "规则：五子连珠",
        "AI：策略-价值网络 + MCTS",
        "对手：纯 MCTS",
        "结果：AI 第 {} 手获胜".format(len(moves)),
    ]
    y = 205
    for line in info:
        draw.text((610, y), line, font=small_font, fill=INK)
        y += 38
    img.save(path)


def build_assets(results, sweep):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    build_graph_logic_image(ASSET_DIR / "gomoku_graph_logic.png")
    build_mcts_image(ASSET_DIR / "mcts_tree_logic.png")
    build_eval_image(ASSET_DIR / "eval_summary.png", results, sweep)
    build_sample_board_image(
        ASSET_DIR / "sample_game_board.png",
        results["matches"]["alpha_vs_pure_mcts"]["sample_game"],
    )


def set_run_font(run, size=10.5, bold=False, color=None, east_asia="微软雅黑"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, fill=None, size=9.2, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align or (WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) <= 12 else WD_ALIGN_PARAGRAPH.LEFT)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)


def set_table_borders(table, color="AAB8C7"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:{}".format(edge)))
        if element is None:
            element = OxmlElement("w:{}".format(edge))
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.3)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Calibri"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return doc


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = "Heading {}".format(level)
    p.paragraph_format.space_before = Pt(6 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=14.2 if level == 1 else 11.8, bold=True, color=BLUE)
    return p


def add_body(doc, text, size=10.3, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=8.8, color=(90, 90, 90))


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="D3DCE8")
    cell = table.cell(0, 0)
    shade_cell(cell, "F5F7FA")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        run.font.size = Pt(8.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths=None, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, fill=LIGHT_BLUE, size=font_size)
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_number_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("AlphaGo 与围棋 AI 家族实验报告")
        set_run_font(run, size=8.5, color=(100, 100, 100))


def pct(value):
    return "{:.0%}".format(value)


def build_report():
    results = json.loads(RESULTS_PATH.read_text(encoding="utf-8"))
    sweep = json.loads(SWEEP_PATH.read_text(encoding="utf-8"))
    build_assets(results, sweep)

    cfg = results["config"]
    random_summary = results["matches"]["alpha_vs_random"]["summary"]
    heuristic_summary = results["matches"]["alpha_vs_heuristic"]["summary"]
    pure_summary = results["matches"]["alpha_vs_pure_mcts"]["summary"]
    first_move = results["first_move_probe"]
    self_play = results["self_play_probe"]

    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(28)
    title.paragraph_format.space_after = Pt(5)
    r = title.add_run("AlphaGo 与围棋 AI 家族")
    set_run_font(r, size=24, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("技术原理调研、8x8 五子棋 AlphaZero 复现与评估")
    set_run_font(r, size=14.5, bold=True, color=(70, 70, 70))

    meta = doc.add_table(rows=5, cols=2)
    set_table_borders(meta, color="CBD6E2")
    meta_rows = [
        ("课程", "人工智能科学与技术"),
        ("小组/成员", "待补充：姓名、学号、班级"),
        ("复现模型", "junxiaosong/AlphaZero_Gomoku：best_policy_8_8_5.model"),
        ("直接演示", "python play_gomoku_8x8.py --playouts 200"),
        ("完成日期", "2026 年 5 月 21 日"),
    ]
    for row, (k, v) in zip(meta.rows, meta_rows):
        set_cell_text(row.cells[0], k, bold=True, fill="EAF2F8", size=9.8)
        set_cell_text(row.cells[1], v, size=9.8)

    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        "本报告围绕 AlphaGo 与围棋 AI 家族展开，重点回答课程要求中的问题定义、研究意义、关键人工智能技术、典型案例、争议挑战与小组观点。"
        "实验部分采用 8x8、五子连珠的 AlphaZero-Gomoku 预训练模型，而不是较小的 6x6 四子棋模型；该模型可直接在人机对弈脚本中使用，也能通过自动评估脚本复现实验数据。",
    )
    add_body(
        doc,
        "报告的核心观点是：AlphaGo 的历史意义不只在“赢棋”，而在于展示了深度学习、强化学习、自我对弈和蒙特卡洛树搜索组合成闭环决策系统的方式。"
        "五子棋复现实验虽然远小于 19x19 围棋，但足以展示“策略-价值网络给出先验，MCTS 在搜索树上校正先验，自我对弈产生训练标签”的关键机制。",
    )
    add_heading(doc, "任务完成对应表", 2)
    add_table(
        doc,
        ["课程要求", "本报告对应内容"],
        [
            ["技术原理调研", "第 2-4 节：围棋 AI 家族、策略-价值网络、MCTS、自我对弈"],
            ["关键代码实现", "第 5 节：8x8 模型加载、MCTS 调用、训练目标和新增可玩脚本"],
            ["运行结果", "第 6 节：三类基线评估、搜索次数敏感性、样例对局"],
            ["小结和分析", "第 7-8 节：争议挑战、小组观点与局限"],
        ],
        widths=[4.5, 12.5],
        font_size=8.8,
    )

    doc.add_page_break()
    add_heading(doc, "1. 问题定义与研究意义", 1)
    add_body(
        doc,
        "围棋 AI 与五子棋 AI 都属于完全信息、零和、离散动作的序贯决策问题。给定棋盘状态 s，智能体要选择动作 a，使长期胜率或最终收益最大化。"
        "难点在于：局部战斗与全局布局相互影响，当前一步的价值往往要经过很多手之后才显现；同时棋盘状态空间巨大，单纯暴力搜索无法覆盖全部变化。",
    )
    add_body(
        doc,
        "AlphaGo 的重要性在于，它把人类长期依赖经验和直觉的围棋决策问题转化为可学习、可搜索、可迭代改进的计算过程。"
        "DeepMind 官方资料将 Go 描述为长期 AI 挑战，Nature 2016 论文报告 AlphaGo 将策略网络、价值网络与树搜索结合，战胜欧洲冠军樊麾并在 2016 年对李世石取得 4:1。"
    )
    add_body(
        doc,
        "本次实验选择五子棋作为复现载体，是因为五子棋保留了棋类博弈中“连线、威胁、阻断、先后手、搜索深度”等核心结构，同时规模足够小，能够在普通电脑上快速运行和展示。"
        "这使报告能够从课堂角度看清 AlphaZero 思想，而不被完整围棋模型的算力成本淹没。",
    )

    add_heading(doc, "2. 围棋 AI 家族与典型案例", 1)
    add_table(
        doc,
        ["系统/事件", "时间", "核心机制", "代表意义"],
        [
            ["AlphaGo", "2015-2016", "人类棋谱监督学习 + 自我对弈强化学习 + 策略/价值网络 + MCTS", "首次在围棋领域击败顶尖职业棋手，成为深度强化学习标志性事件"],
            ["AlphaGo Zero", "2017", "不使用人类棋谱，仅从规则出发自我对弈训练单一策略-价值网络", "证明强棋力并不依赖记忆人类棋谱，关键在自我改进闭环"],
            ["AlphaZero", "2017", "把 AlphaGo Zero 思路推广到国际象棋、日本将棋和围棋", "展示规则驱动自我对弈方法的通用性"],
            ["Leela Zero", "2017 起", "开源复现 AlphaGo Zero 路线，社区分布式训练权重", "降低研究者理解神经网络 + MCTS 的门槛"],
            ["KataGo", "2019 起", "开源围棋引擎，加入目数、地盘、规则等辅助预测", "更贴近复盘、教学和实际围棋分析"],
        ],
        widths=[3.0, 1.8, 6.4, 6.0],
        font_size=8.2,
    )
    add_body(
        doc,
        "2016 年 AlphaGo 与李世石五番棋是最具传播性的案例。第二局第 37 手被广泛视为“非人类直觉”的代表，第四局李世石第 78 手也说明人类仍能通过高创造性局部手段给 AI 制造困难。"
        "这些案例把“AI 是否具有创造性”“AI 是否只是背棋谱”“胜率最大化是否等于人类理解的好棋”等讨论带入公众视野。",
    )

    doc.add_page_break()
    add_heading(doc, "3. 技术原理调研", 1)
    add_body(
        doc,
        "AlphaGo/AlphaZero 路线可以概括为“神经网络给出直觉，树搜索负责深思”。策略网络输出各合法落子的先验概率 p，价值网络输出局面估值 v；MCTS 在搜索树上反复选择、扩展、评估、回传，用访问次数构造更稳健的改进策略 π。",
    )
    add_table(
        doc,
        ["技术环节", "作用", "在五子棋复现中的对应"],
        [
            ["状态表示", "把棋盘转成多通道张量，保留当前方、对方、最近一步、执棋方等信息", "Board.current_state() 返回 4×8×8 张量"],
            ["策略-价值网络", "共享卷积层提取棋形，策略头输出落子分布，价值头输出胜负估计", "policy_value_net_pytorch.py 中三层卷积 + 双头网络"],
            ["MCTS/PUCT", "用 Q+U 在搜索树中平衡利用与探索，网络先验指导扩展", "mcts_alphaZero.py 中 select、expand、update_recursive"],
            ["自我对弈", "当前模型与自己对弈，生成 (s, π, z) 训练样本", "Game.start_self_play() 返回训练样本流"],
            ["数据增强", "利用棋盘旋转/翻转对称性扩大样本，减少过拟合", "train.py 中 get_equi_data() 产生 8 倍样本"],
        ],
        widths=[3.0, 7.0, 6.8],
        font_size=8.2,
    )
    doc.add_picture(str(ASSET_DIR / "gomoku_graph_logic.png"), width=Cm(16.6))
    add_caption(doc, "图 1  五子棋的图论视角：棋盘交叉点是节点，横/竖/斜相邻关系是边，五连胜对应同色路径。")

    add_body(
        doc,
        "从图论角度看，五子棋可建模为棋盘图 G=(V,E)：每个交叉点是节点，横向、纵向和两条斜向的相邻关系构成边。"
        "一个五连就是同色棋子在某一方向上形成长度为 4 的路径；活三、冲四等威胁可以理解为尚未闭合但可扩展的路径结构。"
        "神经网络负责学习这些局部子图模式，MCTS 则评估某条路径扩展后对全局胜率的影响。",
    )

    doc.add_page_break()
    add_heading(doc, "4. MCTS 与强化学习闭环", 1)
    doc.add_picture(str(ASSET_DIR / "mcts_tree_logic.png"), width=Cm(16.6))
    add_caption(doc, "图 2  AlphaZero 风格 MCTS：选择、扩展、评估、回传组成搜索树上的策略改进过程。")
    add_body(
        doc,
        "MCTS 的选择公式可写为 Q(s,a)+U(s,a)。Q 表示该分支历史模拟的平均价值，U 与策略网络先验 P(s,a)、父节点访问次数 N(s) 和子节点访问次数 N(s,a) 有关。"
        "先验概率高但访问少的动作会获得探索奖励；多次模拟后，访问次数分布就成为训练目标 π，比网络原始输出更稳健。",
    )
    add_code_block(
        doc,
        [
            "# mcts_alphaZero.py 的核心选择逻辑",
            "U(s,a) = c_puct * P(s,a) * sqrt(N(s)) / (1 + N(s,a))",
            "a* = argmax_a [ Q(s,a) + U(s,a) ]",
            "",
            "# policy_value_net_pytorch.py 的训练目标",
            "loss = (z - v)^2 - sum_a pi(a|s) * log p(a|s)",
            "# z 是最终胜负，pi 是 MCTS 访问次数归一化后的搜索策略。"
        ],
    )
    add_body(
        doc,
        "训练闭环可表述为：当前网络 fθ(s) 先输出 (p,v)，MCTS 用 (p,v) 进行搜索并得到改进策略 π；对局结束后得到结果 z；训练网络使 p 接近 π、v 接近 z。"
        "因此模型并不是简单模仿已有棋谱，而是在反复自我对弈中把搜索得到的经验压缩回神经网络参数。",
    )

    add_heading(doc, "5. 模型复现与可直接使用版本", 1)
    add_body(
        doc,
        "根据“棋盘稍大、便于演示、必须是五子棋模型”的要求，本次使用仓库自带的 8x8 五子连珠预训练模型 best_policy_8_8_5.model。"
        "相比 6x6 四子棋，它更接近常见五子棋规则；相比 15x15 以上模型，它又能在普通 CPU 上稳定演示。"
    )
    add_table(
        doc,
        ["文件", "用途"],
        [
            ["play_gomoku_8x8.py", "新增：直接人机对弈脚本，默认加载 8x8 五子棋模型，支持 --playouts、--ai-first、hint"],
            ["experiments/run_gomoku_experiment.py", "新增/改造：自动评估 8x8 模型，对随机、启发式、纯 MCTS 三类基线输出 JSON"],
            ["experiments/evaluate_playout_sweep.py", "新增：评估不同搜索次数对耗时与胜率的影响"],
            ["best_policy_8_8_5.model", "仓库自带预训练策略-价值网络权重，用 NumPy 推理路径加载"],
        ],
        widths=[5.5, 11.2],
        font_size=8.5,
    )
    add_code_block(
        doc,
        [
            "# 直接使用：人机对弈",
            "python play_gomoku_8x8.py --playouts 200",
            "",
            "# 让 AI 先手，降低搜索次数以便课堂快速演示",
            "python play_gomoku_8x8.py --ai-first --playouts 80",
            "",
            "# 复现实验评估",
            "python experiments/run_gomoku_experiment.py --games 12 --pure-games 6 \\",
            "    --alpha-playouts 120 --pure-playouts 120 --board-size 8 --n-in-row 5 \\",
            "    --model-file best_policy_8_8_5.model --out experiments/results_8x8.json",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "6. 专业评估设计与运行结果", 1)
    add_body(
        doc,
        "评估不只看能否下一步棋，而是从对手类型、先后手、公平搜索预算、耗时和样例对局几个维度观察模型。"
        "实验环境为 Windows + Codex bundled Python 3.12.13 + NumPy 2.3.5，CPU 推理；AI 每步默认 120 次 MCTS playouts，纯 MCTS 基线也使用 120 次 playouts，尽量保证搜索预算公平。",
    )
    add_table(
        doc,
        ["评估项", "设置", "目的"],
        [
            ["对随机玩家", "12 局，AI 先后手各半", "检验模型是否掌握基本连线和阻断"],
            ["对启发式玩家", "12 局，启发式会立即取胜、阻断、延展最长线", "检验模型是否强于规则型战术基线"],
            ["对纯 MCTS", "6 局，同等 120 playouts", "检验神经网络先验是否提升搜索效率"],
            ["搜索次数敏感性", "40/80/120/200 playouts，对启发式玩家各 6 局", "观察胜率和平均耗时随搜索预算变化"],
        ],
        widths=[3.4, 6.4, 6.8],
        font_size=8.4,
    )
    add_table(
        doc,
        ["对手", "局数", "AI 胜", "对手胜", "和棋", "AI 得分", "先手得分", "后手得分", "平均手数", "平均每步耗时"],
        [
            ["随机", random_summary["games"], random_summary["alpha_wins"], random_summary["opponent_wins"], random_summary["ties"], pct(random_summary["alpha_score"]), pct(random_summary["alpha_score_first"]), pct(random_summary["alpha_score_second"]), random_summary["avg_moves"], "{}s".format(random_summary["avg_sec_per_move"])],
            ["启发式", heuristic_summary["games"], heuristic_summary["alpha_wins"], heuristic_summary["opponent_wins"], heuristic_summary["ties"], pct(heuristic_summary["alpha_score"]), pct(heuristic_summary["alpha_score_first"]), pct(heuristic_summary["alpha_score_second"]), heuristic_summary["avg_moves"], "{}s".format(heuristic_summary["avg_sec_per_move"])],
            ["纯 MCTS", pure_summary["games"], pure_summary["alpha_wins"], pure_summary["opponent_wins"], pure_summary["ties"], pct(pure_summary["alpha_score"]), pct(pure_summary["alpha_score_first"]), pct(pure_summary["alpha_score_second"]), pure_summary["avg_moves"], "{}s".format(pure_summary["avg_sec_per_move"])],
        ],
        widths=[2.2, 1.4, 1.4, 1.5, 1.2, 1.7, 1.8, 1.8, 1.8, 2.6],
        font_size=7.5,
    )
    add_body(
        doc,
        "结果显示，8x8 预训练模型在三类基线上均取得 100% 得分；首步探测选择坐标 {}，位于中心区域；一次自我对弈 {} 手结束，产生 {} 条原始样本，按旋转/翻转增强后可得到 {} 条样本。"
        .format(first_move["selected_location"], self_play["moves"], self_play["raw_samples"], self_play["augmented_samples_if_training"]),
    )
    doc.add_picture(str(ASSET_DIR / "eval_summary.png"), width=Cm(16.6))
    add_caption(doc, "图 3  左：三类对手下的得分；右：搜索次数增加带来更高耗时，本次设置下胜率已饱和。")

    doc.add_page_break()
    add_heading(doc, "7. 结果分析、样例对局与局限", 1)
    doc.add_picture(str(ASSET_DIR / "sample_game_board.png"), width=Cm(13.8))
    add_caption(doc, "图 4  AlphaZero-MCTS 对纯 MCTS 的样例终局。黑棋为 AI，编号表示落子顺序。")
    add_body(
        doc,
        "样例对局中，AI 首步落在中心附近，随后沿斜向不断制造多方向威胁，最终形成五连。这个过程体现了神经网络先验与 MCTS 的互补：网络快速筛出高价值区域，搜索再验证不同防守和进攻分支。"
        "纯 MCTS 虽然有相同 playouts，但缺少学习到的先验，更多依赖随机 rollout，搜索资源更容易浪费在低质量分支上。",
    )
    add_table(
        doc,
        ["playouts/步", "对启发式玩家局数", "AI 得分", "平均手数", "平均每步耗时"],
        [
            [s["alpha_playouts"], s["summary"]["games"], pct(s["summary"]["alpha_score"]), s["summary"]["avg_moves"], "{}s".format(s["summary"]["avg_sec_per_move"])]
            for s in sweep["settings"]
        ],
        widths=[3.0, 3.5, 2.4, 2.4, 3.6],
        font_size=8.6,
    )
    add_body(
        doc,
        "搜索次数敏感性实验表明，从 40 到 200 playouts，模型对启发式玩家的得分均为 100%，但平均每步耗时从 0.05s 增至 0.20s。"
        "这说明在课堂演示时可以将 --playouts 设为 80 或 120，兼顾响应速度和稳定性；若用于更强对手，可提高到 200 以上。",
    )
    add_body(
        doc,
        "局限也必须说明：本实验使用的是 8x8 小棋盘，不等价于完整 15x15 或 19x19 棋盘；评估局数有限，胜率只能说明该演示模型强于设定基线，不能代表所有五子棋程序；预训练权重来自开源仓库，本报告重点是复现、评估和解释，而非从零训练一个大规模模型。",
    )

    add_heading(doc, "8. 争议挑战与小组观点", 1)
    add_body(
        doc,
        "围绕 AlphaGo 的常见误解是“AI 只是背棋谱”。AlphaGo Zero 的结果反驳了这一点：只给规则和胜负目标，系统也能通过自我对弈形成强策略。"
        "更准确的说法是，AI 将搜索得到的有效经验压缩进网络参数，而不是逐条存储人类棋谱。",
    )
    add_body(
        doc,
        "另一个争议是创造性。第 37 手等案例看起来像创造性，但其来源是目标函数、训练分布和搜索机制；它能产生人类少见的高胜率策略，却不等同于人类意义上的意图表达。"
        "此外，围棋 AI 通常最大化胜率而不是赢更多目，因此一些“缓手”或保守选择可能是风险控制，而不是模型不会下。",
    )
    add_body(
        doc,
        "我们的小组观点是：AlphaGo 的核心价值不是“机器打败人类”，而是证明了一种通用智能构造范式：用神经网络学习复杂状态的压缩表示，用搜索做可校验规划，再用自我对弈把经验反哺模型。"
        "这种范式把感知、判断、规划和学习连接成闭环，是后来许多强化学习、规划和自动决策系统的重要思想来源。",
    )
    add_body(
        doc,
        "本次 8x8 五子棋复现适合课堂展示：它既能直接运行、让同学现场对弈，又能通过评估表和搜索树解释技术机制。若后续继续扩展，可安装 PyTorch 从随机初始化训练若干轮，记录 loss、KL、胜率曲线；也可接入简单 GUI，增强 PPT 展示效果。",
    )

    add_heading(doc, "参考资料", 1)
    refs = [
        ("Google DeepMind：AlphaGo 项目页", "https://deepmind.google/research/alphago/"),
        ("Silver et al. Mastering the game of Go with deep neural networks and tree search. Nature, 2016", "https://www.nature.com/articles/nature16961"),
        ("Silver et al. Mastering the game of Go without human knowledge. Nature, 2017", "https://www.nature.com/articles/nature24270"),
        ("Silver et al. Mastering Chess and Shogi by Self-Play with a General Reinforcement Learning Algorithm. arXiv, 2017", "https://arxiv.org/abs/1712.01815"),
        ("junxiaosong/AlphaZero_Gomoku GitHub 仓库", "https://github.com/junxiaosong/AlphaZero_Gomoku"),
        ("Leela Zero GitHub 仓库", "https://github.com/leela-zero/leela-zero"),
        ("KataGo Distributed Training", "https://katagotraining.org/"),
        ("GFW Jumper：聊聊大伙儿对 AlphaGo 的误解", "https://gfwjumper.medium.com/%E8%81%8A%E8%81%8A%E5%A4%A7%E4%BC%99%E5%84%BF-%E5%8C%85%E6%8B%AC%E6%9F%90%E4%BA%9B%E8%81%8C%E4%B8%9A%E5%9B%B4%E6%A3%8B%E6%89%8B-%E5%AF%B9-alphago-%E7%9A%84%E8%AF%AF%E8%A7%A3-e6fd5d034123"),
    ]
    for i, (label, url) in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run("[{}] {}. ".format(i, label))
        set_run_font(run, size=8.1)
        add_hyperlink(p, url, url)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    run = note.add_run("需人工补充：封面中的小组成员姓名、学号、班级。")
    set_run_font(run, size=9.0, bold=True, color=(192, 80, 77))

    add_page_number_footer(doc)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_report())
