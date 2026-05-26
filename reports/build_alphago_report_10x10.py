# -*- coding: utf-8 -*-
"""Build the upgraded 10x10 Gomoku report."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
PROJECT = ROOT / "AlphaZero_Gomoku"
RESULTS = PROJECT / "experiments" / "results_10x10.json"
TRAIN_LOG = PROJECT / "experiments" / "training_log_10x10.json"
OUT = ROOT / "AlphaGo与围棋AI家族_实验报告_最终充实版.docx"
ASSETS = ROOT / "reports" / "assets_10x10"
FIGURE_FONT_PATH = Path("C:/Windows/Fonts/simhei.ttf")
BODY_FONT_NAME = "宋体"
HEADING_FONT_NAME = "黑体"
LATIN_BODY_FONT_NAME = "Times New Roman"
LATIN_HEADING_FONT_NAME = "Arial"

BLUE = (31, 78, 121)
RED = (192, 80, 77)
INK = (35, 35, 35)
PAPER = (250, 252, 255)


def font(size):
    return ImageFont.truetype(str(FIGURE_FONT_PATH), size) if FIGURE_FONT_PATH.exists() else ImageFont.load_default()


def center_text(draw, box, text, fnt, fill=INK):
    l, t, r, b = box
    bb = draw.textbbox((0, 0), text, font=fnt)
    draw.text((l + (r - l - bb[2] + bb[0]) / 2, t + (b - t - bb[3] + bb[1]) / 2), text, font=fnt, fill=fill)


def draw_centered_lines(draw, box, lines, fnt, fill=INK, line_gap=5):
    l, t, r, b = box
    heights = []
    widths = []
    for line in lines:
        bb = draw.textbbox((0, 0), line, font=fnt)
        widths.append(bb[2] - bb[0])
        heights.append(bb[3] - bb[1])
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = t + (b - t - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        draw.text((l + (r - l - w) / 2, y), line, font=fnt, fill=fill)
        y += h + line_gap


def draw_wrapped_text(draw, xy, text, fnt, fill=INK, max_width=360, line_gap=9):
    x, y = xy
    current = ""
    for ch in text:
        candidate = current + ch
        bb = draw.textbbox((0, 0), candidate, font=fnt)
        if current and bb[2] - bb[0] > max_width:
            draw.text((x, y), current, font=fnt, fill=fill)
            y += (bb[3] - bb[1]) + line_gap
            current = ch
        else:
            current = candidate
    if current:
        draw.text((x, y), current, font=fnt, fill=fill)
        bb = draw.textbbox((0, 0), current, font=fnt)
        y += (bb[3] - bb[1]) + line_gap
    return y


def make_coordinate_image(path):
    img = Image.new("RGB", (1200, 580), PAPER)
    d = ImageDraw.Draw(img)
    d.text((40, 30), "演示界面坐标修正：10x10 坐标从 1 开始", font=font(34), fill=BLUE)
    x0, y0, cell = 150, 130, 46
    for c in range(10):
        center_text(d, (x0 + c * cell - 20, y0 - 42, x0 + c * cell + 20, y0 - 8), str(c + 1), font(18))
    for r in range(10):
        center_text(d, (x0 - 60, y0 + r * cell - 18, x0 - 25, y0 + r * cell + 18), str(r + 1), font(18))
    for r in range(10):
        d.line((x0, y0 + r * cell, x0 + 9 * cell, y0 + r * cell), fill=(145, 155, 165), width=2)
        d.line((x0 + r * cell, y0, x0 + r * cell, y0 + 9 * cell), fill=(145, 155, 165), width=2)
    for r in range(10):
        for c in range(10):
            x, y = x0 + c * cell, y0 + r * cell
            d.ellipse((x - 4, y - 4, x + 4, y + 4), fill=(92, 102, 112))
    moves = [(6, 5, "O"), (5, 5, "X"), (6, 6, "O"), (4, 6, "X"), (7, 7, "O")]
    for row, col, piece in moves:
        x, y = x0 + (col - 1) * cell, y0 + (row - 1) * cell
        fill = (34, 34, 34) if piece == "O" else (245, 245, 245)
        outline = (33, 96, 170)
        d.ellipse((x - 19, y - 19, x + 19, y + 19), fill=fill, outline=outline, width=3)
    d.rounded_rectangle((700, 145, 1110, 440), radius=18, fill=(236, 243, 250), outline=(176, 198, 220), width=2)
    notes = [
        "输入 5,5：第 5 行第 5 列",
        "第 1 行在棋盘最上方",
        "第 10 列在棋盘最右侧",
    ]
    y = 195
    for note in notes:
        d.text((735, y), note, font=font(22), fill=INK)
        y += 68
    img.save(path)


def make_graph_image(path):
    img = Image.new("RGB", (1200, 620), PAPER)
    d = ImageDraw.Draw(img)
    d.text((40, 28), "图论意味：五子棋是棋盘图上的路径搜索问题", font=font(34), fill=BLUE)
    x0, y0, cell = 140, 125, 42
    for i in range(10):
        d.line((x0, y0 + i * cell, x0 + 9 * cell, y0 + i * cell), fill=(145, 155, 165), width=2)
        d.line((x0 + i * cell, y0, x0 + i * cell, y0 + 9 * cell), fill=(145, 155, 165), width=2)
    for r in range(10):
        for c in range(10):
            x, y = x0 + c * cell, y0 + r * cell
            d.ellipse((x - 4, y - 4, x + 4, y + 4), fill=(92, 102, 112))
    path_nodes = [(5, 4), (6, 5), (7, 6), (8, 7), (9, 8)]
    for a, b in zip(path_nodes, path_nodes[1:]):
        ax, ay = x0 + (a[1] - 1) * cell, y0 + (a[0] - 1) * cell
        bx, by = x0 + (b[1] - 1) * cell, y0 + (b[0] - 1) * cell
        d.line((ax, ay, bx, by), fill=(33, 96, 170), width=6)
    for row, col in path_nodes:
        x, y = x0 + (col - 1) * cell, y0 + (row - 1) * cell
        d.ellipse((x - 17, y - 17, x + 17, y + 17), fill=(35, 35, 35), outline=(33, 96, 170), width=3)
    d.rounded_rectangle((650, 140, 1120, 485), radius=18, fill=(236, 243, 250), outline=(176, 198, 220), width=2)
    lines = [
        "节点 V：棋盘交叉点",
        "边 E：横、竖、斜相邻关系",
        "五连：5 个同色节点形成连续路径",
        "威胁：开放路径、冲四、活三",
        "MCTS：搜索不同路径扩展后的胜率",
    ]
    y = 170
    for line in lines:
        y = draw_wrapped_text(d, (680, y), line, font(23), fill=INK, max_width=395, line_gap=7) + 22
    img.save(path)


def make_mcts_image(path):
    img = Image.new("RGB", (1200, 640), PAPER)
    d = ImageDraw.Draw(img)
    d.text((40, 28), "MCTS 与强化学习闭环：搜索改进策略，策略再指导搜索", font=font(32), fill=BLUE)
    nodes = {
        "s0": (600, 125, "s0\nN=80"),
        "a": (260, 305, "a1\nP=0.22\nQ=0.31"),
        "b": (600, 305, "a2\nP=0.46\nQ=0.58"),
        "c": (940, 305, "a3\nP=0.09\nQ=0.10"),
        "d": (490, 485, "leaf\nv=0.71"),
        "e": (710, 485, "leaf\nv=-0.16"),
    }
    for s, t in [("s0", "a"), ("s0", "b"), ("s0", "c"), ("b", "d"), ("b", "e")]:
        sx, sy, _ = nodes[s]
        tx, ty, _ = nodes[t]
        d.line((sx, sy + 54, tx, ty - 54), fill=(125, 145, 165), width=4)
    for name, (x, y, text) in nodes.items():
        fill = (211, 232, 214) if name == "b" else (224, 238, 251)
        d.rounded_rectangle((x - 92, y - 54, x + 92, y + 54), radius=15, fill=fill, outline=(70, 100, 130), width=3)
        draw_centered_lines(d, (x - 82, y - 45, x + 82, y + 45), text.split("\n"), font(17), line_gap=8)
    d.rounded_rectangle((90, 555, 1110, 610), radius=14, fill=(245, 248, 252), outline=(190, 205, 220), width=2)
    d.text((115, 572), "搜索思想：用 Q 表示历史平均价值，用 U 奖励高先验、低访问次数的分支。", font=font(20), fill=INK)
    img.save(path)


def make_training_image(path, training):
    img = Image.new("RGB", (1200, 620), PAPER)
    d = ImageDraw.Draw(img)
    d.text((40, 28), "10x10 自我对弈训练日志：样本累计与损失变化", font=font(34), fill=BLUE)
    games = training["games"]
    x0, y0, w, h = 95, 130, 990, 360
    d.line((x0, y0 + h, x0 + w, y0 + h), fill=(80, 80, 80), width=2)
    d.line((x0, y0, x0, y0 + h), fill=(80, 80, 80), width=2)
    max_buffer = max(g["buffer"] for g in games)
    max_loss = max(g["policy_loss"] for g in games)
    points_buffer = []
    points_loss = []
    for i, g in enumerate(games):
        x = x0 + int(i / max(1, len(games) - 1) * w)
        yb = y0 + h - int(g["buffer"] / max_buffer * (h - 30))
        yl = y0 + h - int(g["policy_loss"] / max_loss * (h - 30))
        points_buffer.append((x, yb))
        points_loss.append((x, yl))
    for pts, color in [(points_buffer, (33, 96, 170)), (points_loss, RED)]:
        for a, b in zip(pts, pts[1:]):
            d.line((*a, *b), fill=color, width=4)
        for x, y in pts[::4]:
            d.ellipse((x - 5, y - 5, x + 5, y + 5), fill=color)
    d.text((120, 510), "蓝线：经验池样本数（最终 {} 条）".format(max_buffer), font=font(20), fill=(33, 96, 170))
    d.text((520, 510), "红线：策略交叉熵损失（随自我对弈目标变化而波动）", font=font(20), fill=RED)
    d.text((120, 550), "训练：{} 局自我对弈，{} 次 MCTS/步，耗时 {} 秒".format(training["config"]["self_play_games"], training["config"]["self_play_playouts"], training["elapsed_sec"]), font=font(18), fill=(80, 80, 80))
    img.save(path)


def make_eval_image(path, results):
    img = Image.new("RGB", (1200, 620), PAPER)
    d = ImageDraw.Draw(img)
    d.text((40, 28), "10x10 强化训练模型评估：基线胜率与搜索预算", font=font(34), fill=BLUE)
    summaries = [
        ("随机", results["matches"]["ai_vs_random"]["summary"]),
        ("战术启发式", results["matches"]["ai_vs_tactical"]["summary"]),
    ]
    x0, y0, bottom = 110, 155, 500
    d.line((x0, bottom, 520, bottom), fill=(90, 90, 90), width=2)
    d.line((x0, y0, x0, bottom), fill=(90, 90, 90), width=2)
    for i in range(6):
        y = bottom - i * 60
        d.line((x0, y, 520, y), fill=(225, 230, 235), width=1)
        d.text((55, y - 12), "{:.1f}".format(i / 5), font=font(16), fill=(90, 90, 90))
    for i, (name, summary) in enumerate(summaries):
        x = x0 + 95 + i * 145
        bar_h = int(summary["score"] * 300)
        d.rectangle((x, bottom - bar_h, x + 85, bottom), fill=(33, 96, 170))
        d.text((x + 5, bottom - bar_h - 34), "{:.1%}".format(summary["score"]), font=font(20), fill=BLUE)
        center_text(d, (x - 40, bottom + 15, x + 125, bottom + 45), name, font(18))
    d.text((130, 110), "正式评估得分", font=font(22), fill=INK)
    x1 = 690
    d.text((x1, 110), "搜索次数敏感性（对战术启发式）", font=font(22), fill=INK)
    d.line((x1, bottom, 1110, bottom), fill=(90, 90, 90), width=2)
    d.line((x1, y0, x1, bottom), fill=(90, 90, 90), width=2)
    sweep = results["playout_sweep"]
    max_t = max(s["summary"]["avg_sec_per_move"] for s in sweep)
    pts = []
    for i, s in enumerate(sweep):
        px = x1 + 65 + i * 135
        py = bottom - int(s["summary"]["avg_sec_per_move"] / max_t * 260)
        pts.append((px, py))
        d.ellipse((px - 8, py - 8, px + 8, py + 8), fill=RED)
        d.text((px - 30, py - 34), "{}s".format(s["summary"]["avg_sec_per_move"]), font=font(17), fill=RED)
        center_text(d, (px - 45, bottom + 15, px + 45, bottom + 45), str(s["playouts"]), font(18))
    for a, b in zip(pts, pts[1:]):
        d.line((*a, *b), fill=RED, width=4)
    d.text((x1 + 15, 545), "横轴：playouts/步；纵向：平均每步耗时。小样本下胜率有波动。", font=font(17), fill=(80, 80, 80))
    img.save(path)


def build_assets(results, training):
    ASSETS.mkdir(parents=True, exist_ok=True)
    make_coordinate_image(ASSETS / "coordinate_fix.png")
    make_graph_image(ASSETS / "graph_logic.png")
    make_mcts_image(ASSETS / "mcts_logic.png")
    make_training_image(ASSETS / "training_curve.png", training)
    make_eval_image(ASSETS / "eval_10x10.png", results)


def set_font(run, size=10.5, bold=False, color=None, east_asia=BODY_FONT_NAME, latin=None):
    if latin is None:
        latin = LATIN_BODY_FONT_NAME
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def borders(table, color="AAB8C7"):
    tbl_pr = table._tbl.tblPr
    b = tbl_pr.first_child_found_in("w:tblBorders")
    if b is None:
        b = OxmlElement("w:tblBorders")
        tbl_pr.append(b)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = b.find(qn("w:{}".format(edge)))
        if e is None:
            e = OxmlElement("w:{}".format(edge))
            b.append(e)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)


def cell_text(cell, text, header=False, size=8.6):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) < 12 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(str(text))
    set_font(r, size=size, bold=header, east_asia=HEADING_FONT_NAME if header else BODY_FONT_NAME)
    if header:
        shade(cell, "DDEBF7")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None, size=8.8):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    borders(t)
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, header=True, size=size)
        if widths:
            t.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, size=size)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = "Heading {}".format(level)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=15, bold=True, color=BLUE, east_asia=HEADING_FONT_NAME, latin=LATIN_HEADING_FONT_NAME)


def add_body(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_font(r, size=size, east_asia=BODY_FONT_NAME)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_font(r, size=9.0, color=(90, 90, 90), east_asia=BODY_FONT_NAME)


def add_code(doc, lines):
    t = doc.add_table(rows=1, cols=1)
    borders(t, color="D3DCE8")
    c = t.cell(0, 0)
    shade(c, "F5F7FA")
    c.text = ""
    p = c.paragraphs[0]
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        r.font.size = Pt(8.2)


def add_formula_table(doc):
    add_table(
        doc,
        ["项目", "表达式", "含义"],
        [
            ["探索项 U", "U(s,a)=c_puct·P(s,a)·sqrt(N(s))/(1+N(s,a))", "先验概率越高、访问次数越少，探索奖励越大"],
            ["动作选择", "a*=argmax_a[Q(s,a)+U(s,a)]", "在平均价值 Q 和探索项 U 之间折中"],
            ["策略损失", "policy_loss=-Σπ(a|s)log p(a|s)", "让模型预测接近 MCTS 访问次数分布"],
            ["价值损失", "value_loss=(z-v)^2", "让局面价值接近最终胜负结果"],
        ],
        widths=[2.5, 6.2, 7.9],
        size=8.4,
    )


def add_command_table(doc):
    add_table(
        doc,
        ["场景", "运行方式", "说明"],
        [
            ["直接演示", "python play_gomoku_10x10.py --playouts 80", "人类先手，适合课堂现场对弈"],
            ["AI 先手", "python play_gomoku_10x10.py --ai-first --playouts 80", "展示模型首步和中心控制倾向"],
            ["重新训练", "python train_lightweight_10x10.py --games 24 --playouts 20", "复现实验中的 10x10 自我对弈训练"],
        ],
        widths=[2.4, 7.5, 6.7],
        size=8.4,
    )


def hyperlink(paragraph, text, url):
    rid = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), rid)
    r = OxmlElement("w:r")
    rp = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), LATIN_BODY_FONT_NAME)
    rf.set(qn("w:hAnsi"), LATIN_BODY_FONT_NAME)
    rf.set(qn("w:eastAsia"), BODY_FONT_NAME)
    rp.append(rf)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "17")
    rp.append(sz)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), "0563C1")
    rp.append(col)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rp.append(u)
    r.append(rp)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    h.append(r)
    paragraph._p.append(h)


def pct(x):
    return "{:.1%}".format(float(x))


def build():
    results = json.loads(RESULTS.read_text(encoding="utf-8"))
    training = json.loads(TRAIN_LOG.read_text(encoding="utf-8"))
    build_assets(results, training)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.55)
    sec.bottom_margin = Cm(1.45)
    sec.left_margin = Cm(1.65)
    sec.right_margin = Cm(1.65)
    doc.styles["Normal"].font.name = LATIN_BODY_FONT_NAME
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_NAME)
    doc.styles["Normal"].font.size = Pt(10.5)
    for style_name in ("Heading 1", "Heading 2", "Title", "Subtitle"):
        style = doc.styles[style_name]
        style.font.name = LATIN_HEADING_FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT_NAME)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(25)
    r = p.add_run("AlphaGo 与围棋 AI 家族")
    set_font(r, size=24, bold=True, color=BLUE, east_asia=HEADING_FONT_NAME, latin=LATIN_HEADING_FONT_NAME)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("技术调研、10x10 五子棋强化训练与可演示模型")
    set_font(r, size=14.5, bold=True, color=(70, 70, 70), east_asia=HEADING_FONT_NAME, latin=LATIN_HEADING_FONT_NAME)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["课程", "人工智能科学与技术"],
            ["小组/成员", "待补充：姓名、学号、班级"],
            ["最终演示模型", "10x10 轻量策略-价值模型（模型文件见 models/）"],
            ["直接运行", "python play_gomoku_10x10.py --playouts 80"],
            ["完成日期", "2026 年 5 月 22 日"],
        ],
        widths=[4.2, 12.5],
        size=9.2,
    )
    add_heading(doc, "摘要", 1)
    add_body(doc, "本报告围绕 AlphaGo 与围棋 AI 家族展开，说明问题定义、研究意义、关键技术、典型案例、争议挑战与小组观点。根据演示需求，本次不再停留在 8x8 预训练模型，而是新增 10x10 五子棋轻量策略-价值模型，并通过自我对弈进行训练。")
    add_body(doc, "实验保留 AlphaZero 的核心接口：策略-价值函数输出动作先验和局面价值，MCTS 在搜索树上改进策略，自我对弈生成训练标签。由于当前环境未安装深度学习框架，10x10 版本采用可训练的图模式特征模型代替 CNN，使训练、评估和人机对弈都能在普通 CPU 上完成。")
    add_table(
        doc,
        ["要求", "对应内容"],
        [
            ["棋盘与演示", "10x10，行号从上到下 1-10，列号从左到右 1-10，修复原先坐标不对齐问题"],
            ["模型训练", "24 局自我对弈、451 条训练样本，保存 lightweight_policy_10x10.json"],
            ["专业评估", "随机基线、战术启发式基线、搜索次数敏感性、训练日志分析"],
            ["图文展示", "棋盘坐标图、图论路径图、MCTS 搜索树图、训练曲线和评估图"],
        ],
        widths=[4.0, 12.8],
        size=8.8,
    )

    doc.add_page_break()
    add_heading(doc, "1. 问题定义与研究意义", 1)
    add_body(doc, "围棋和五子棋都属于完全信息、零和、离散动作的序贯决策问题。给定棋盘局面，智能体需要在合法落子中选择最优动作，使长期胜率或最终收益最大化。难点在于局部战斗和全局布局相互影响，当前一步的价值往往要经过许多手之后才显现。")
    add_body(doc, "从人工智能角度看，这类棋类问题的核心不是“把棋谱存下来”，而是在巨大状态空间中进行近似规划。围棋 19x19 棋盘有极高的分支因子，五子棋虽然规则更简单，但同样存在长程依赖、攻防转换和局部威胁组合。本报告用 10x10 五子棋作为可运行复现实验，是因为它保留了棋盘图搜索、策略评估和自我对弈训练等关键结构，同时训练成本更适合课堂演示。")
    add_body(doc, "AlphaGo 的意义不只是战胜人类棋手，而是把长期依赖经验和直觉的棋类决策转化为可学习、可搜索、可迭代改进的计算过程。DeepMind 官方资料把 Go 称为长期 AI 挑战；Nature 2016 论文将策略网络、价值网络和树搜索结合起来，推动了深度强化学习在复杂决策任务中的应用。")
    add_table(
        doc,
        ["建模要素", "棋类问题中的含义", "本实验中的对应"],
        [
            ["状态", "当前棋盘、执棋方、最近落子等信息", "10x10 Board 记录棋盘与落子历史"],
            ["动作", "在空交叉点上落下一子", "屏幕坐标采用行、列均从 1 开始"],
            ["策略", "对各合法落子的优先级分布", "MCTS 访问次数形成训练目标"],
            ["价值", "当前局面对当前方的胜负倾向", "终局胜负回传为价值标签"],
        ],
        widths=[2.8, 6.7, 7.1],
        size=8.6,
    )
    add_heading(doc, "2. 围棋 AI 家族与典型案例", 1)
    add_body(doc, "围棋 AI 的发展可以看作三条线索的汇合：第一是用神经网络学习人类棋谱中的局面模式，第二是用强化学习通过自我对弈超越已有经验，第三是用蒙特卡洛树搜索把模型判断转化为可验证的局面推演。后续系统在是否使用人类棋谱、是否开源、是否加入规则和目数辅助预测等方面形成了不同分支。")
    add_table(
        doc,
        ["系统/事件", "时间", "核心机制", "代表意义"],
        [
            ["AlphaGo", "2015-2016", "监督学习 + 自我对弈强化学习 + 策略/价值网络 + MCTS", "战胜樊麾、李世石，成为深度强化学习里程碑"],
            ["AlphaGo Zero", "2017", "只从规则出发，自我对弈训练单一策略-价值网络", "证明强棋力不依赖记忆人类棋谱"],
            ["AlphaZero", "2017", "把 AlphaGo Zero 思路推广到国际象棋、日本将棋和围棋", "展示规则驱动自我对弈方法的通用性"],
            ["Leela Zero", "2017 起", "开源复现 AlphaGo Zero 路线，社区分布式训练", "降低理解神经网络 + MCTS 的门槛"],
            ["KataGo", "2019 起", "加入目数、地盘、规则等辅助预测", "更适合复盘、教学和实战分析"],
        ],
        widths=[3.0, 1.8, 6.5, 5.8],
        size=7.8,
    )
    add_body(doc, "李世石五番棋第二局第 37 手常被视为“非人类直觉”的代表，第四局李世石第 78 手也说明人类仍可通过高创造性局部手段制造困难。这些案例让“AI 是否只是背棋谱”“AI 是否有创造性”“胜率最大化是否等于人类理解的好棋”等问题进入公众讨论，也为本实验的复现部分提供了参照：我们关注的不是单步下法多么神奇，而是模型、搜索和训练闭环如何共同产生决策。")

    doc.add_page_break()
    add_heading(doc, "3. 技术原理：策略-价值网络、MCTS 与图结构", 1)
    add_body(doc, "这一路线可以概括为“网络给出直觉，搜索负责深思”。策略函数给出每个合法落子的先验概率，价值函数估计当前局面对当前方的胜负价值；MCTS 反复选择、扩展、评估和回传，用访问次数得到更稳健的搜索策略。")
    add_body(doc, "在 AlphaGo/AlphaZero 框架中，策略输出不是最终答案，而是搜索的先验；价值输出也不是绝对真理，而是对叶节点局面的快速估计。两者结合后，系统既能避免纯搜索的巨大计算量，也能避免单纯神经网络“一步判断”缺乏推演的问题。这种结构体现了学习与规划的结合，是棋类 AI 能从规则走向高水平决策的关键。")
    add_table(
        doc,
        ["技术环节", "作用", "本次 10x10 复现方式"],
        [
            ["状态表示", "从当前方视角描述棋盘、最近落子和执棋方", "Board 保存棋盘；轻量模型提取路径、开放端、中心等特征"],
            ["策略-价值模型", "输出落子先验 p 和局面价值 v", "LightweightPolicyValueNet，可训练特征权重"],
            ["MCTS/PUCT", "用 Q+U 在搜索树中兼顾利用和探索", "沿用 mcts_alphaZero.py 的选择、扩展、回传逻辑"],
            ["自我对弈训练", "用当前模型生成 (s, π, z) 样本再更新模型", "train_lightweight_10x10.py 完成 24 局训练"],
        ],
        widths=[3.0, 6.5, 7.2],
        size=8.3,
    )
    add_body(doc, "本次 10x10 复现没有使用深度卷积网络，而是用可训练的图模式特征模型近似策略-价值函数。模型显式提取连续棋子长度、开放端数量、立即成五、阻断对方成五、中心控制和近邻密度等特征；这些特征对应五子棋中的基本战术概念，适合在无 GPU、无深度学习框架的环境中完成可解释复现。")
    doc.add_picture(str(ASSETS / "graph_logic.png"), width=Cm(16.6))
    add_caption(doc, "图 1  图论视角：棋盘交叉点是节点，横/竖/斜相邻关系是边，五连胜对应同色路径。")
    add_body(doc, "从图论角度看，五子棋可建模为棋盘图 G=(V,E)。一个五连就是同色棋子在某一方向上形成 5 个节点、4 条边的连续路径；活三、冲四等威胁则是尚未闭合但可扩展的路径结构。模型学习这些局部图模式，MCTS 评估路径扩展后的全局胜率。")

    doc.add_page_break()
    add_heading(doc, "4. MCTS 与训练闭环", 1)
    doc.add_picture(str(ASSETS / "mcts_logic.png"), width=Cm(16.6))
    add_caption(doc, "图 2  MCTS 在搜索树上用 Q+U 选择分支，并把叶节点价值回传到祖先节点。")
    add_body(doc, "MCTS 的选择公式可写成 Q(s,a)+U(s,a)。Q 表示该动作历史模拟的平均价值，U 与策略先验 P(s,a)、父节点访问次数 N(s) 和子节点访问次数 N(s,a) 有关。先验高但访问少的动作会获得探索奖励，多次模拟后的访问次数分布就是训练目标 π。")
    add_body(doc, "一次搜索通常包含四个阶段：选择阶段沿着 Q+U 最大的边向下走；扩展阶段为新局面加入合法子节点；评估阶段调用策略-价值模型得到先验概率和局面价值；回传阶段把价值沿路径反向更新到祖先节点。重复多次后，根节点的访问次数分布比单次网络输出更稳定，因此可以作为更强的改进策略。")
    add_formula_table(doc)
    add_body(doc, "强化学习闭环的关键在于“搜索改进模型，模型再指导搜索”。自我对弈时，当前模型先通过 MCTS 产生更强的落子分布；对局结束后，胜负结果作为价值标签；训练脚本再用策略损失和价值损失调整权重。这样即使没有人工棋谱，模型也能从自己生成的经验中迭代。")
    doc.add_page_break()
    add_heading(doc, "5. 10x10 演示脚本与训练实现", 1)
    add_body(doc, "演示脚本已修复坐标问题：行号从上到下 1-10，列号从左到右 1-10，输入也采用同一套屏幕坐标。这样课堂演示时不需要解释内部数组坐标，观众可以直接按棋盘上的数字落子。")
    add_body(doc, "实现上，演示脚本把“屏幕坐标”和“程序内部坐标”分离：用户输入的第 r 行第 c 列先被转换为 Board 使用的位置编号，再交给 MCTS 玩家搜索。这样既能保持原仓库 Board 类的接口不变，又能让展示界面符合直觉，避免原先 0-7 标签和棋盘点不对齐的问题。")
    doc.add_picture(str(ASSETS / "coordinate_fix.png"), width=Cm(16.6))
    add_caption(doc, "图 3  演示界面坐标修正示意：第 1 行在最上方，第 10 列在最右侧。")
    add_table(
        doc,
        ["文件", "作用"],
        [
            ["play_gomoku_10x10.py", "最终演示脚本；支持 --playouts、--ai-first、hint"],
            ["lightweight_policy_value_10x10.py", "可训练策略-价值模型，提取五连、阻断、开放端、中心控制等特征"],
            ["train_lightweight_10x10.py", "自我对弈训练脚本，保存模型和训练日志"],
            ["experiments/evaluate_10x10_model.py", "正式评估脚本，输出 results_10x10.json"],
        ],
        widths=[5.2, 11.5],
        size=8.5,
    )
    add_body(doc, "轻量策略-价值模型对每个候选落子计算一组局部图特征，并用可训练权重得到先验分数；价值头则从双方最长连续线、开放端、威胁数量、中心控制等信息估计局面倾向。虽然它不是大规模 CNN，但接口形式与 AlphaZero 模型一致，因此可以无缝接入 mcts_alphaZero.py。")
    add_command_table(doc)

    doc.add_page_break()
    add_heading(doc, "6. 训练过程与专业评估", 1)
    cfg = training["config"]
    total_samples = sum(g["samples"] for g in training["games"])
    add_body(doc, "训练设置为 10x10 棋盘、五子连珠，进行 {} 局自我对弈，每步 {} 次 MCTS playouts，共产生 {} 条训练样本，最终经验池保留 {} 条样本，耗时 {} 秒。".format(cfg["self_play_games"], cfg["self_play_playouts"], total_samples, training["games"][-1]["buffer"], training["elapsed_sec"]))
    add_body(doc, "每条训练样本包含三个部分：当前局面、MCTS 在该局面下形成的访问次数分布、最终胜负结果。策略部分学习“搜索后更认可哪些落子”，价值部分学习“该局面最终更接近胜还是负”。这种训练目标比直接模仿某一步落子更稳定，因为它把搜索中的多分支评估压缩进监督信号。")
    doc.add_picture(str(ASSETS / "training_curve.png"), width=Cm(16.6))
    add_caption(doc, "图 4  训练日志：经验池逐步扩大，策略损失随自我对弈目标变化而波动。")
    add_table(
        doc,
        ["阶段", "随机基线得分", "战术启发式得分", "说明"],
        [
            ["训练前", pct(training["before_eval"]["random"]["score"]), pct(training["before_eval"]["tactical"]["score"]), "初始特征权重已包含基本棋理"],
            ["训练后", pct(training["after_eval"]["random"]["score"]), pct(training["after_eval"]["tactical"]["score"]), "自我对弈后中心、近邻、连线权重发生调整"],
            ["正式评估", pct(results["matches"]["ai_vs_random"]["summary"]["score"]), pct(results["matches"]["ai_vs_tactical"]["summary"]["score"]), "使用独立评估脚本、35 playouts/步"],
        ],
        widths=[3.0, 3.0, 3.4, 7.0],
        size=8.5,
    )
    add_body(doc, "评估时设置了随机基线和战术启发式基线。随机基线用于验证模型是否学到基本合法下法和直接胜负判断；战术启发式基线会优先考虑成五、阻断、中心和连线，能更好检验模型是否具备实际攻防能力。需要说明的是，本次轻量模型不是 CNN，而是为了当前环境可训练、可演示而设计的策略-价值近似器。它仍然遵循 AlphaZero 的训练接口：MCTS 产生 π，终局产生 z，模型通过策略损失和价值损失更新权重。")

    doc.add_page_break()
    add_heading(doc, "7. 运行结果与分析", 1)
    rsum = results["matches"]["ai_vs_random"]["summary"]
    tsum = results["matches"]["ai_vs_tactical"]["summary"]
    add_table(
        doc,
        ["对手", "局数", "AI 胜", "对手胜", "和棋", "得分", "先手得分", "后手得分", "平均手数", "平均每步耗时"],
        [
            ["随机", rsum["games"], rsum["ai_wins"], rsum["opponent_wins"], rsum["ties"], pct(rsum["score"]), pct(rsum["score_first"]), pct(rsum["score_second"]), rsum["avg_moves"], "{}s".format(rsum["avg_sec_per_move"])],
            ["战术启发式", tsum["games"], tsum["ai_wins"], tsum["opponent_wins"], tsum["ties"], pct(tsum["score"]), pct(tsum["score_first"]), pct(tsum["score_second"]), tsum["avg_moves"], "{}s".format(tsum["avg_sec_per_move"])],
        ],
        widths=[2.4, 1.2, 1.3, 1.5, 1.2, 1.6, 1.8, 1.8, 1.8, 2.5],
        size=7.7,
    )
    add_body(doc, "正式评估中，模型对随机玩家取得 100.0% 得分，对战术启发式玩家取得 87.5% 得分。首步探测选择屏幕坐标 {}，位于 10x10 棋盘中心区域，符合五子棋中中心控制与多方向延展的直觉。".format(results["first_move_probe"]["selected_display_location"]))
    add_body(doc, "这里的得分按胜局计 1 分、和棋计 0.5 分、负局计 0 分计算，并分别记录先手和后手表现。区分先后手是必要的，因为五子棋存在明显的先手优势；如果只看总胜率，可能会掩盖模型在后手防守、阻断威胁和延迟失败方面的能力。")
    doc.add_picture(str(ASSETS / "eval_10x10.png"), width=Cm(16.6))
    add_caption(doc, "图 5  左：正式评估得分；右：搜索次数增加会提高耗时，小样本下胜率有波动。")
    add_table(
        doc,
        ["playouts/步", "局数", "得分", "平均手数", "平均每步耗时"],
        [[s["playouts"], s["summary"]["games"], pct(s["summary"]["score"]), s["summary"]["avg_moves"], "{}s".format(s["summary"]["avg_sec_per_move"])] for s in results["playout_sweep"]],
        widths=[3.0, 2.5, 2.5, 3.0, 3.8],
        size=8.5,
    )
    add_body(doc, "搜索次数敏感性实验显示，10x10 棋盘比 8x8 更耗时；playouts 增加时平均每步耗时明显上升。由于评估局数较少，胜率曲线并不单调，这并不意味着更多搜索一定更差，而是说明小样本评估会受到先后手、开局随机性和对手策略波动影响。更严格的实验应扩大局数、固定随机种子并报告置信区间。课堂演示建议使用 --playouts 50 到 80，以保证响应速度。")

    doc.add_page_break()
    add_heading(doc, "8. 争议挑战、小组观点与局限", 1)
    add_body(doc, "围绕 AlphaGo 的常见误解是“AI 只是背棋谱”。事实上，只给规则和胜负目标，系统也能通过自我对弈形成强策略。更准确的说法是，AI 把搜索得到的有效经验压缩进模型参数，而不是逐条记忆人类棋谱。")
    add_body(doc, "另一个争议是创造性。李世石对局中的第 37 手等案例看起来像创造性，但其来源是目标函数、训练分布和搜索机制；它能产生人类少见的高胜率策略，却不等同于人类意义上的意图表达。围棋 AI 通常最大化胜率而不是赢更多目，因此某些保守选择可能是风险控制。")
    add_body(doc, "我们的小组观点是：这类系统的核心价值不是“机器打败人类”，而是展示了一种通用智能构造方式：用模型学习复杂状态的压缩表示，用搜索做可校验规划，再用自我对弈把经验反哺模型。本次 10x10 五子棋模型虽然轻量，但完整演示了这一闭环。")
    add_body(doc, "局限也必须明确：10x10 版本为了可训练和可演示，使用的是图模式特征模型而非大规模深度卷积网络；训练局数有限，评估局数也有限，因此结论只说明它强于本报告设置的基线，不能代表完整 15x15 五子棋或 19x19 围棋的棋力。")
    add_table(
        doc,
        ["问题", "争议焦点", "本报告观点"],
        [
            ["可解释性", "神经网络和搜索树给出的胜率不等同于人类语言解释", "用图模式、访问次数和基线实验补充解释，但不能把胜率直接当作因果说明"],
            ["泛化能力", "小棋盘复现能否代表围棋或标准五子棋", "本实验用于说明技术闭环，不宣称达到完整棋类系统棋力"],
            ["评估可靠性", "少量对局容易受先后手和随机性影响", "应扩大局数、固定随机种子、分别报告先手和后手表现"],
        ],
        widths=[2.5, 6.1, 8.0],
        size=7.9,
    )
    add_heading(doc, "参考资料", 1)
    refs = [
        ("Google DeepMind：AlphaGo 项目页", "https://deepmind.google/research/alphago/"),
        ("Silver et al. Mastering the game of Go with deep neural networks and tree search. Nature, 2016", "https://www.nature.com/articles/nature16961"),
        ("Silver et al. Mastering the game of Go without human knowledge. Nature, 2017", "https://www.nature.com/articles/nature24270"),
        ("Silver et al. Mastering Chess and Shogi by Self-Play with a General Reinforcement Learning Algorithm. arXiv, 2017", "https://arxiv.org/abs/1712.01815"),
        ("junxiaosong/AlphaZero_Gomoku GitHub 仓库", "https://github.com/junxiaosong/AlphaZero_Gomoku"),
        ("Leela Zero GitHub 仓库", "https://github.com/leela-zero/leela-zero"),
        ("KataGo Distributed Training", "https://katagotraining.org/"),
        ("GFW Jumper：聊聊大伙儿对 AlphaGo 的误解", "https://gfwjumper.medium.com/%E8%81%8A%E8%81%8A%E5%A4%A7%E4%BC%99%E5%84%BF-%E5%8C%85%E6%8B%AC%E6%9F%90%E4%BA%9B%E8%81%8C%E4%B8%9A%E5%9B%B4%E6%A3%8B%E6%89%8B-%E5%AF%B9-alphago-%E7%9A%84%E8%AF%AF%E8%A7%A3-e6fd5d034123"),
    ]
    for i, (label, url) in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run("[{}] {}. ".format(i, label))
        set_font(r, size=8.2)
        hyperlink(p, url, url)
    p = doc.add_paragraph()
    r = p.add_run("需人工补充：封面中的小组成员姓名、学号、班级。")
    set_font(r, size=9.0, bold=True, color=RED)

    for section in doc.sections:
        f = section.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = f.add_run("AlphaGo 与围棋 AI 家族实验报告")
        set_font(r, size=8.5, color=(100, 100, 100))
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
